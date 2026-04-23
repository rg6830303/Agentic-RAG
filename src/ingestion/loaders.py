from __future__ import annotations

import csv
import json
import re
from pathlib import Path
from typing import Callable

import pandas as pd

from src.utils.hashing import checksum_file, checksum_text
from src.utils.models import DocumentSection, LoadedDocument
from src.utils.text import normalize_text
from src.utils.time import utc_now_iso


TEXT_EXTENSIONS = {
    ".txt",
    ".md",
    ".py",
    ".js",
    ".ts",
    ".tsx",
    ".jsx",
    ".java",
    ".cpp",
    ".c",
    ".cs",
    ".go",
    ".rs",
    ".html",
    ".css",
    ".json",
    ".yaml",
    ".yml",
    ".xml",
    ".sql",
    ".csv",
}
SUPPORTED_EXTENSIONS = TEXT_EXTENSIONS | {".pdf", ".docx"}


def is_probably_binary(path: Path, sample_size: int = 1024) -> bool:
    try:
        chunk = path.read_bytes()[:sample_size]
    except OSError:
        return True
    if not chunk:
        return False
    return b"\x00" in chunk


class IngestionError(RuntimeError):
    """Raised when a file cannot be ingested."""


class DocumentLoader:
    def __init__(self) -> None:
        self._custom_loaders: dict[str, Callable[[Path], list[DocumentSection]]] = {
            ".pdf": self._load_pdf,
            ".docx": self._load_docx,
            ".csv": self._load_csv,
            ".json": self._load_json,
        }

    def load(self, path: Path) -> LoadedDocument:
        extension = path.suffix.lower()
        if extension not in SUPPORTED_EXTENSIONS:
            raise IngestionError(f"Unsupported extension: {extension}")
        if extension in TEXT_EXTENSIONS and is_probably_binary(path):
            raise IngestionError(f"Skipping binary-looking file: {path.name}")
        loader = self._custom_loaders.get(extension, self._load_text)
        sections = loader(path)
        if not sections:
            raise IngestionError(f"No text could be extracted from {path.name}")
        checksum = checksum_file(path)
        document_id = checksum_text(f"{path.resolve()}::{checksum}")[:24]
        return LoadedDocument(
            document_id=document_id,
            file_path=str(path.resolve()),
            file_name=path.name,
            extension=extension,
            checksum=checksum,
            ingested_at=utc_now_iso(),
            sections=sections,
            metadata={
                "section_count": len(sections),
                "supported_extension": extension in SUPPORTED_EXTENSIONS,
            },
        )

    def _load_text(self, path: Path) -> list[DocumentSection]:
        raw_text = path.read_text(encoding="utf-8", errors="ignore")
        normalized = normalize_text(raw_text)
        if not normalized:
            return []
        return [
            DocumentSection(
                section_id=checksum_text(f"{path.name}:0")[:24],
                text=normalized,
            )
        ]

    def _load_csv(self, path: Path) -> list[DocumentSection]:
        try:
            frame = pd.read_csv(path)
            text = frame.to_csv(index=False)
        except Exception:
            with path.open("r", encoding="utf-8", errors="ignore", newline="") as handle:
                reader = csv.reader(handle)
                rows = [", ".join(cell.strip() for cell in row) for row in reader]
            text = "\n".join(rows)
        normalized = normalize_text(text)
        return [
            DocumentSection(
                section_id=checksum_text(f"{path.name}:csv")[:24],
                text=normalized,
                heading="csv_table",
            )
        ]

    def _load_json(self, path: Path) -> list[DocumentSection]:
        raw = path.read_text(encoding="utf-8", errors="ignore")
        try:
            parsed = json.loads(raw)
            pretty = json.dumps(parsed, indent=2, ensure_ascii=True)
        except json.JSONDecodeError:
            pretty = raw
        normalized = normalize_text(pretty)
        return [
            DocumentSection(
                section_id=checksum_text(f"{path.name}:json")[:24],
                text=normalized,
                heading="json_document",
            )
        ]

    def _load_pdf(self, path: Path) -> list[DocumentSection]:
        try:
            from pypdf import PdfReader
        except ImportError as exc:
            raise IngestionError(
                "PDF ingestion requires the optional dependency 'pypdf'."
            ) from exc
        reader = PdfReader(str(path))
        sections: list[DocumentSection] = []
        for index, page in enumerate(reader.pages, start=1):
            extracted = normalize_text(page.extract_text() or "")
            if extracted:
                sections.append(
                    DocumentSection(
                        section_id=checksum_text(f"{path.name}:pdf:{index}")[:24],
                        text=extracted,
                        page_number=index,
                        heading=f"Page {index}",
                    )
                )
        return sections

    def _load_docx(self, path: Path) -> list[DocumentSection]:
        try:
            from docx import Document
        except ImportError as exc:
            raise IngestionError(
                "DOCX ingestion requires the optional dependency 'python-docx'."
            ) from exc
        document = Document(str(path))
        blocks: list[str] = []
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                blocks.append(text)
        combined = normalize_text("\n\n".join(blocks))
        if not combined:
            return []
        headings = re.findall(r"(?m)^(.*:)$", combined)
        return [
            DocumentSection(
                section_id=checksum_text(f"{path.name}:docx")[:24],
                text=combined,
                heading=headings[0] if headings else None,
            )
        ]
